import os
import json
import sqlite3
import uuid
import secrets
from datetime import datetime
from functools import wraps
from flask import Flask, request, jsonify, render_template, g, send_file, session, redirect, url_for
from werkzeug.security import generate_password_hash, check_password_hash
from PyPDF2 import PdfReader
from docx import Document
from openai import OpenAI

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(BASE_DIR, 'config.json')
DB_FILE = os.path.join(BASE_DIR, 'data', 'papers.db')
UPLOAD_DIR = os.path.join(BASE_DIR, 'uploads')

os.makedirs(os.path.dirname(DB_FILE), exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)


def load_config():
    # Environment variables override file config (for cloud deployment)
    config = {}
    try:
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            config = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        pass
    # Env vars take priority
    if os.environ.get('DEEPSEEK_API_KEY'):
        config['deepseek_api_key'] = os.environ['DEEPSEEK_API_KEY']
    if os.environ.get('DEEPSEEK_BASE_URL'):
        config['deepseek_base_url'] = os.environ['DEEPSEEK_BASE_URL']
    if os.environ.get('DEEPSEEK_MODEL'):
        config['model'] = os.environ['DEEPSEEK_MODEL']
    return config


def save_config(data):
    config = load_config()
    config.update(data)
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


# Persistent secret key
_secret = os.environ.get('SECRET_KEY', '') or load_config().get('secret_key', '') or secrets.token_hex(32)
if not os.environ.get('SECRET_KEY') and not load_config().get('secret_key'):
    save_config({'secret_key': _secret})
app.secret_key = _secret


def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(DB_FILE)
        g.db.row_factory = sqlite3.Row
    return g.db


@app.teardown_appcontext
def close_db(exception):
    db = g.pop('db', None)
    if db is not None:
        db.close()


def init_db():
    db = sqlite3.connect(DB_FILE)
    db.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    db.execute('''
        CREATE TABLE IF NOT EXISTS papers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL DEFAULT 1,
            filename TEXT NOT NULL,
            file_path TEXT NOT NULL,
            title TEXT DEFAULT '',
            content_text TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            share_token TEXT DEFAULT '',
            page_count INTEGER DEFAULT 0,
            file_size INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    ''')
    cols = [r[1] for r in db.execute("PRAGMA table_info(papers)").fetchall()]
    if 'notes' not in cols:
        db.execute("ALTER TABLE papers ADD COLUMN notes TEXT DEFAULT ''")
    if 'user_id' not in cols:
        db.execute("ALTER TABLE papers ADD COLUMN user_id INTEGER NOT NULL DEFAULT 1")
    if 'share_token' not in cols:
        db.execute("ALTER TABLE papers ADD COLUMN share_token TEXT DEFAULT ''")
    db.execute('''
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            paper_id INTEGER NOT NULL,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (paper_id) REFERENCES papers(id) ON DELETE CASCADE
        )
    ''')
    db.commit()
    db.close()


# --- Auth helpers ---
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            if request.is_json or request.path.startswith('/api/'):
                return jsonify({'error': '请先登录'}), 401
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated


def get_current_user():
    if 'user_id' not in session:
        return None
    db = get_db()
    return db.execute("SELECT id, username FROM users WHERE id = ?",
                      (session['user_id'],)).fetchone()


# --- Text extraction ---
def extract_pdf_text(filepath):
    reader = PdfReader(filepath)
    raw_lines = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            raw_lines.append(t)
    raw = '\n'.join(raw_lines)
    return reflow_pdf_text(raw), len(reader.pages)


def extract_docx_text(filepath):
    doc = Document(filepath)
    paras = [para.text for para in doc.paragraphs if para.text.strip()]
    return '\n\n'.join(paras), len(doc.paragraphs)


def reflow_pdf_text(raw_text):
    lines = raw_text.split('\n')
    if not lines:
        return ''
    SENTENCE_END = set('.。?？!！')
    result = []
    buf = ''
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buf:
                result.append(buf)
                buf = ''
            result.append('')
            continue
        if buf and buf[-1] in SENTENCE_END:
            result.append(buf)
            buf = stripped
        elif buf:
            buf += ' ' + stripped
        else:
            buf = stripped
    if buf:
        result.append(buf)
    paragraphs = []
    for block in result:
        if not block:
            if paragraphs and paragraphs[-1]:
                paragraphs.append('')
        else:
            paragraphs.append(block)
    return '\n\n'.join(paragraphs)


# --- AI ---
def get_deepseek_client():
    config = load_config()
    api_key = config.get('deepseek_api_key', '')
    base_url = config.get('deepseek_base_url', 'https://api.deepseek.com')
    if not api_key:
        return None
    return OpenAI(api_key=api_key, base_url=base_url)


def ask_ai(paper_text, question):
    client = get_deepseek_client()
    if not client:
        return "请先在设置中配置 DeepSeek API Key。"
    max_text_len = 12000
    if len(paper_text) > max_text_len:
        paper_text = paper_text[:max_text_len] + "\n\n[文本过长，已截断...]"
    config = load_config()
    model = config.get('model', 'deepseek-chat')
    if paper_text.strip():
        system_prompt = f"""你是一个智能助手，可以自由回答用户的任何问题。

当前用户选中了一篇论文，以下是论文内容供你参考（如果问题与论文相关请结合内容回答；否则忽略论文直接回答）：

## 参考论文内容
{paper_text}"""
    else:
        system_prompt = "你是一个智能助手，请自由回答用户的任何问题。"
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            temperature=0.7,
            max_tokens=4096,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI 调用失败: {str(e)}"


# ==================== Auth Routes ====================
@app.route('/login')
def login_page():
    if 'user_id' in session:
        return redirect(url_for('index'))
    return render_template('login.html')


@app.route('/register')
def register_page():
    if 'user_id' in session:
        return redirect(url_for('index'))
    return render_template('register.html')


@app.route('/api/auth/register', methods=['POST'])
def api_register():
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not username or not password:
        return jsonify({'error': '用户名和密码不能为空'}), 400
    if len(username) < 3:
        return jsonify({'error': '用户名至少3个字符'}), 400
    if len(password) < 4:
        return jsonify({'error': '密码至少4个字符'}), 400
    db = get_db()
    existing = db.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
    if existing:
        return jsonify({'error': '用户名已存在'}), 400
    db.execute("INSERT INTO users (username, password_hash) VALUES (?, ?)",
               (username, generate_password_hash(password)))
    db.commit()
    return jsonify({'ok': True})


@app.route('/api/auth/login', methods=['POST'])
def api_login():
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    if not username or not password:
        return jsonify({'error': '用户名和密码不能为空'}), 400
    db = get_db()
    user = db.execute("SELECT * FROM users WHERE username = ?", (username,)).fetchone()
    if not user or not check_password_hash(user['password_hash'], password):
        return jsonify({'error': '用户名或密码错误'}), 401
    session['user_id'] = user['id']
    session['username'] = user['username']
    return jsonify({'ok': True, 'username': user['username']})


@app.route('/api/auth/logout', methods=['POST'])
def api_logout():
    session.clear()
    return jsonify({'ok': True})


@app.route('/api/auth/me')
def api_me():
    user = get_current_user()
    if not user:
        return jsonify({'logged_in': False})
    return jsonify({'logged_in': True, 'username': user['username']})


# ==================== Main Page ====================
@app.route('/')
@login_required
def index():
    return render_template('index.html')



# ==================== Share Route (no login needed) ====================
@app.route('/s/<token>')
def shared_paper(token):
    db = get_db()
    paper = db.execute("SELECT * FROM papers WHERE share_token = ?", (token,)).fetchone()
    if not paper:
        return "论文不存在或链接已失效", 404
    return render_template('share.html', paper=dict(paper))


@app.route('/api/shared-file/<token>')
def api_shared_file(token):
    db = get_db()
    paper = db.execute("SELECT file_path, filename FROM papers WHERE share_token = ?",
                       (token,)).fetchone()
    if not paper:
        return jsonify({'error': '文件不存在'}), 404
    return send_file(paper['file_path'], download_name=paper['filename'])


# ==================== API Routes (login required) ====================
@app.route('/api/papers', methods=['GET'])
@login_required
def api_list_papers():
    db = get_db()
    user_id = session['user_id']
    query = request.args.get('q', '').strip()
    if query:
        papers = db.execute(
            "SELECT id, filename, title, page_count, file_size, share_token, created_at FROM papers WHERE user_id = ? AND (title LIKE ? OR filename LIKE ?) ORDER BY created_at DESC",
            (user_id, f'%{query}%', f'%{query}%')
        ).fetchall()
    else:
        papers = db.execute(
            "SELECT id, filename, title, page_count, file_size, share_token, created_at FROM papers WHERE user_id = ? ORDER BY created_at DESC",
            (user_id,)
        ).fetchall()
    return jsonify([dict(p) for p in papers])


@app.route('/api/papers/<int:paper_id>', methods=['GET'])
@login_required
def api_get_paper(paper_id):
    db = get_db()
    paper = db.execute("SELECT * FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    return jsonify(dict(paper))


@app.route('/api/papers/<int:paper_id>', methods=['PUT'])
@login_required
def api_update_paper(paper_id):
    db = get_db()
    paper = db.execute("SELECT id FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    data = request.json
    if 'title' in data:
        db.execute("UPDATE papers SET title = ? WHERE id = ?", (data['title'].strip(), paper_id))
        db.commit()
    return jsonify({'ok': True})


@app.route('/api/papers/<int:paper_id>/file')
@login_required
def api_get_paper_file(paper_id):
    db = get_db()
    paper = db.execute("SELECT file_path, filename FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    return send_file(paper['file_path'], download_name=paper['filename'])


@app.route('/api/papers', methods=['POST'])
@login_required
def api_upload_paper():
    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '未选择文件'}), 400
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.pdf', '.docx', '.doc'):
        return jsonify({'error': '仅支持 PDF 和 Word (.pdf/.docx/.doc) 文件'}), 400
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = f"{ts}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, safe_name)
    file.save(file_path)
    file_size = os.path.getsize(file_path)
    content_text = ''
    page_count = 0
    try:
        if ext == '.pdf':
            content_text, page_count = extract_pdf_text(file_path)
        else:
            content_text, page_count = extract_docx_text(file_path)
    except Exception as e:
        os.remove(file_path)
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 400
    title = os.path.splitext(file.filename)[0]
    db = get_db()
    cursor = db.execute(
        "INSERT INTO papers (user_id, filename, file_path, title, content_text, page_count, file_size) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (session['user_id'], file.filename, file_path, title, content_text, page_count, file_size)
    )
    db.commit()
    paper_id = cursor.lastrowid
    paper = db.execute("SELECT * FROM papers WHERE id = ?", (paper_id,)).fetchone()
    return jsonify(dict(paper)), 201


@app.route('/api/papers/<int:paper_id>', methods=['DELETE'])
@login_required
def api_delete_paper(paper_id):
    db = get_db()
    paper = db.execute("SELECT file_path FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    try:
        os.remove(paper['file_path'])
    except FileNotFoundError:
        pass
    db.execute("DELETE FROM papers WHERE id = ?", (paper_id,))
    db.commit()
    return jsonify({'ok': True})


@app.route('/api/papers/<int:paper_id>/share', methods=['POST'])
@login_required
def api_toggle_share(paper_id):
    db = get_db()
    paper = db.execute("SELECT * FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    if paper['share_token']:
        db.execute("UPDATE papers SET share_token = '' WHERE id = ?", (paper_id,))
        db.commit()
        return jsonify({'shared': False, 'url': ''})
    else:
        token = uuid.uuid4().hex[:12]
        db.execute("UPDATE papers SET share_token = ? WHERE id = ?", (token, paper_id))
        db.commit()
        return jsonify({'shared': True, 'url': f'/s/{token}'})


@app.route('/api/ask', methods=['POST'])
def api_ask_general():
    data = request.json
    question = data.get('question', '').strip()
    if not question:
        return jsonify({'error': '问题不能为空'}), 400
    answer = ask_ai('', question)
    return jsonify({'question': question, 'answer': answer, 'created_at': datetime.now().isoformat()})


@app.route('/api/papers/<int:paper_id>/ask', methods=['POST'])
@login_required
def api_ask_ai(paper_id):
    db = get_db()
    paper = db.execute("SELECT * FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    data = request.json
    question = data.get('question', '').strip()
    if not question:
        return jsonify({'error': '问题不能为空'}), 400
    answer = ask_ai(paper['content_text'], question)
    db.execute(
        "INSERT INTO conversations (paper_id, question, answer) VALUES (?, ?, ?)",
        (paper_id, question, answer)
    )
    db.commit()
    return jsonify({'question': question, 'answer': answer, 'created_at': datetime.now().isoformat()})


@app.route('/api/papers/<int:paper_id>/conversations', methods=['GET'])
@login_required
def api_get_conversations(paper_id):
    db = get_db()
    convs = db.execute(
        "SELECT * FROM conversations WHERE paper_id = ? ORDER BY created_at ASC",
        (paper_id,)
    ).fetchall()
    return jsonify([dict(c) for c in convs])


@app.route('/api/papers/<int:paper_id>/notes', methods=['GET'])
@login_required
def api_get_notes(paper_id):
    db = get_db()
    paper = db.execute("SELECT notes FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    return jsonify({'notes': paper['notes'] or ''})


@app.route('/api/papers/<int:paper_id>/notes', methods=['PUT'])
@login_required
def api_save_notes(paper_id):
    db = get_db()
    paper = db.execute("SELECT id FROM papers WHERE id = ? AND user_id = ?",
                       (paper_id, session['user_id'])).fetchone()
    if not paper:
        return jsonify({'error': '论文不存在'}), 404
    data = request.json
    db.execute("UPDATE papers SET notes = ? WHERE id = ?", (data.get('notes', ''), paper_id))
    db.commit()
    return jsonify({'ok': True})


# Keep API config/model-info routes (no auth needed for checking status)
@app.route('/api/config', methods=['GET'])
def api_get_config():
    config = load_config()
    return jsonify({
        'has_api_key': bool(config.get('deepseek_api_key', '')),
        'base_url': config.get('deepseek_base_url', 'https://api.deepseek.com'),
        'model': config.get('model', 'deepseek-chat')
    })


@app.route('/api/config', methods=['POST'])
@login_required
def api_save_config():
    data = request.json
    save_config({
        'deepseek_api_key': data.get('api_key', ''),
        'deepseek_base_url': data.get('base_url', 'https://api.deepseek.com'),
        'model': data.get('model', 'deepseek-chat')
    })
    return jsonify({'ok': True})


@app.route('/api/model-info', methods=['GET'])
def api_model_info():
    config = load_config()
    api_key = config.get('deepseek_api_key', '')
    base_url = config.get('deepseek_base_url', 'https://api.deepseek.com')
    model = config.get('model', 'deepseek-chat')
    info = {
        'configured': bool(api_key),
        'base_url': base_url,
        'model': model,
        'status': 'unknown'
    }
    if not api_key:
        info['status'] = 'no_key'
        return jsonify(info)
    try:
        client = OpenAI(api_key=api_key, base_url=base_url)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": "hi"}],
            max_tokens=10,
        )
        info['status'] = 'connected'
        info['response_model'] = response.model
        info['usage'] = {
            'prompt_tokens': response.usage.prompt_tokens if response.usage else 0,
            'completion_tokens': response.usage.completion_tokens if response.usage else 0,
        }
    except Exception as e:
        info['status'] = 'error'
        info['error'] = str(e)
    return jsonify(info)


# Initialize DB on import (for Vercel) or on direct run
init_db()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
